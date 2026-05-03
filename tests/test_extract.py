"""Tests for document text extraction."""
import pytest
from pathlib import Path
from backend.services.extract_service import extract_text


class TestExtractDocx:
    def test_extract_text_plain_docx(self, tmp_path):
        """Extract text from a DOCX with only text paragraphs."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Hello World")
        doc.add_paragraph("Second paragraph")
        fpath = tmp_path / "test.docx"
        doc.save(str(fpath))

        result = extract_text(fpath, "docx")
        assert "Hello World" in result
        assert "Second paragraph" in result

    def test_extract_docx_empty(self, tmp_path):
        """Extract from empty DOCX returns empty string."""
        from docx import Document

        doc = Document()
        fpath = tmp_path / "empty.docx"
        doc.save(str(fpath))

        result = extract_text(fpath, "docx")
        assert result == "" or result is None

    def test_extract_docx_with_image_placeholder(self, tmp_path):
        """DOCX with inline image includes image index marker."""
        from docx import Document
        from docx.shared import Inches
        from PIL import Image
        import io

        # Create a small test image
        img = Image.new("RGB", (10, 10), color="red")
        img_bytes = io.BytesIO()
        img.save(img_bytes, format="PNG")
        img_bytes.seek(0)

        doc = Document()
        doc.add_paragraph("Before image")
        doc.add_picture(img_bytes, width=Inches(1))
        doc.add_paragraph("After image")
        fpath = tmp_path / "with_image.docx"
        doc.save(str(fpath))

        result = extract_text(fpath, "docx")
        assert "Before image" in result
        assert "After image" in result
        assert "[image:0]" in result

    def test_extract_docx_saves_images_to_disk(self, tmp_path):
        """DOCX extraction with user_id/doc_id saves images to disk."""
        from docx import Document
        from docx.shared import Inches
        from PIL import Image
        import io
        from unittest.mock import patch

        img = Image.new("RGB", (10, 10), color="blue")
        img_bytes = io.BytesIO()
        img.save(img_bytes, format="PNG")
        img_bytes.seek(0)

        doc = Document()
        doc.add_paragraph("Text with image")
        doc.add_picture(img_bytes, width=Inches(1))
        fpath = tmp_path / "with_image.docx"
        doc.save(str(fpath))

        saved_images = []
        with patch("backend.services.file_service.save_docx_images") as mock_save:
            mock_save.return_value = 1
            result = extract_text(fpath, "docx", user_id=1, doc_id=42)
            assert mock_save.called
            saved_images = mock_save.call_args[0][2]  # third arg is the images list

        assert "[image:0]" in result
        assert len(saved_images) == 1
        assert saved_images[0][0] == ".png"  # extension
        assert len(saved_images[0][1]) > 0    # image bytes

    def test_extract_unknown_type_returns_none(self, tmp_path):
        """Unknown file type returns None."""
        fpath = tmp_path / "test.xyz"
        fpath.write_text("hello")
        result = extract_text(fpath, "xyz")
        assert result is None


class TestExtractMarkdown:
    def test_extract_markdown(self, tmp_path):
        """Extract text from markdown file."""
        fpath = tmp_path / "test.md"
        fpath.write_text("# Title\n\nSome content", encoding="utf-8")
        result = extract_text(fpath, "md")
        assert "Title" in result
        assert "Some content" in result

    def test_extract_markdown_empty(self, tmp_path):
        """Extract from empty markdown returns empty string."""
        fpath = tmp_path / "empty.md"
        fpath.write_text("", encoding="utf-8")
        result = extract_text(fpath, "md")
        assert result == ""
